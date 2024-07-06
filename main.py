from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
import os 
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import base64
from langchain_community.llms import Ollama

load_dotenv()

os.environ['SERPER_API_KEY'] = os.getenv('SERPER_API_KEY')

def generate_docx(result):
    doc = Document()
    doc.add_heading("Healthcare Diagnosis and Treatment Recommendations", 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64decode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

st.set_page_config(
    layout='wide'
)

st.title("Medical AI Bot")

gender = st.selectbox('Select Gender', ("Male", "Female", "Other"))
age = st.number_input("Enter Age", min_value=0, max_value=120, value=25)
symptoms = st.text_area("Enter Symptoms", 'e.g., Fever, Cough, Headache')
medical_history = st.text_area("Enter your medical history",'e.g., Diabetes, Hypertension')


search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = Ollama(
    model = "phi3"
)


#Defining Agents

diagonostician = Agent(
    role ="Medical Diagonostician"
    goal = "Analyze patient symptoms and medical history to provide a perliminary diagnosis.",
    backstory = "This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanded algorithms and medical knowledge to identify potential health issues",
    verbose = True,
    allow_delegation = False
    tools = [search_tool, scrape_tool],
    llm = llm

)

treatment_advisior = Agent(
    role = "Treatment Advisior",
    goal = "Recommend appropriate treatment plans based on the diagnosis provided by the medical Diagnostician",
    backstory="This  agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effectice treatments.",
    verbose = True,
    allow_delegation = False,
    tools = [search_tool, scrape_tool],
    llm = llm
)


# tasks

diagnose_task = Task(
    description = (
        "1. Analyze the patient's symptoms ({symptoms}) and medical history({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output = "A preliminary diadnosis with a list of pssible conditions.",
    agent = diagonostician
)

treatment_task = Task(
    description =(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_histoy}) and current symptoms ({symptoms}).\n "
        "3. Provide detailed treatment recommendations, inculding medicatins, lifestyle changes, and follow-up care."
    ),
    expected_output= " A comprehesive treatment plan tailored to the patient's needs",
    agent = treatment_advisior
)

crew = Crew(
    agents = [ diagonostician, treatment_advisior],
    tasks=[diagnose_task, treatment_task],
    verbose= 2
)


if st.button(" Get Diagnosis and Treatment Plan"):
    with st.spinner("Generating Recommendations....."):
        results = crew.kickoff(inputs={"symptoms":symptoms, "medical_history": medical_history})
        st.write(results)
        docx_file = generate_docx(results)

        download_link = get_download_link(docx_file, "diagnosis_and_treatment_paln.docx")
        st.markdown(download_link, unsafe_allow_html=True)


